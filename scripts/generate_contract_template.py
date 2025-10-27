from docx import Document

# Minimal script to create a DOCX template with placeholders for docxtpl
# This will create/overwrite contract_template.docx in the workspace root

def create_template(path):
    doc = Document()
    doc.add_heading('CONTRATO DE ARRENDAMENTO URBANO', level=1)
    doc.add_paragraph('Para fins Habitacionais')
    doc.add_paragraph('\nENTRE\n')
    doc.add_paragraph('{{ senhorio }}')
    doc.add_paragraph('E')
    doc.add_paragraph('{{ inquilino }}')
    doc.add_paragraph('\n2025\n')
    doc.add_paragraph('CONTRATO DE ARRENDAMENTO URBANO\nPara fins Habitacionais')
    doc.add_paragraph('\nENTRE:\n')
    doc.add_paragraph('{{ senhorio }}, sociedade registada ao abrigo das leis Angolana, com sede social na Rua {{ senhorio_address }} com número contribuinte n.º {{ senhorio_nif }}, representada neste acto pelo Sr. {{ representative_name }}, na qualidade de Gerente, com poderes para o acto, doravante designado, abreviadamente "SENHORIO".')
    doc.add_paragraph('\nE\n')
    doc.add_paragraph('{{ inquilino }}, pessoa singular, portador do {{ document_type }} {{ document_number }}, emitido em {{ document_issue_date }}, válido até {{ document_expiry_date }}, NIF {{ inquilino_nif }}, adiante abreviadamente designado por ARRENDATÁRIO.')
    doc.add_paragraph('\nIndividualmente designadas por “Parte” e colectivamente por “Partes”')
    # Insert a few placeholders used later in the contract
    doc.add_paragraph('\nCLÁUSULA SEGUNDA (Vigência)')
    doc.add_paragraph('O contrato terá o prazo de 1(um) ano com início à {{ start_date_written }} à {{ end_date_written }}, sendo automaticamente renovável por períodos sucessivos de iguais períodos.')
    doc.add_paragraph('\nCLÁUSULA TERCEIRA (Renda e Formas de Pagamento)')
    doc.add_paragraph('As Partes acordam um valor mensal da renda devida pela utilização da fracção é o equivalente AOA {{ valor_renda }} de kwanzas), mensal, ...')
    doc.add_paragraph('O pagamento da renda será efectuado {{ forma_pagamento }} o equivalente a AOA {{ valor_renda }} (--------- de kwanzas), nos primeiros 8 dias úteis.')
    doc.add_paragraph('Sobre o valor mensal da renda ainda terá de ser pago o montante equivalente á AOA {{ valor_caucao }} (----------de kwanzas), correspondente à caução...')
    doc.add_paragraph('A taxa de condomínio a data presente é equivalente à AOA {{ taxa_condominio }} (----------------------------- kwanzas),...')
    doc.add_paragraph('\nCLÁUSULA DÉCIMA TERCEIRA (Notificações)')
    doc.add_paragraph('Os contactos do Arrendatário:')
    doc.add_paragraph('Tel.: {{ inquilino_contact }}')
    doc.add_paragraph('Email: {{ inquilino_email }}')

    doc.save(path)

if __name__ == '__main__':
    create_template('contract_template.docx')
    print('contract_template.docx created/overwritten')
