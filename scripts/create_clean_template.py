from docx import Document

CONTRACT_TEXT = [
    ('heading', 'CONTRATO DE ARRENDAMENTO URBANO'),
    ('para', 'Para fins Habitacionais'),
    ('line', '\nENTRE\n'),
    ('line', '{{ senhorio }}'),
    ('line', 'E'),
    ('line', '{{ inquilino }}'),
    ('line', '\nCONTRATO DE ARRENDAMENTO URBANO\nPara fins Habitacionais'),
    ('line', '\nENTRE:\n'),
    ('para', '{{ senhorio }}, sociedade registada ao abrigo das leis Angolana, com sede social na Rua {{ senhorio_address }} com número contribuinte n.º {{ senhorio_nif }}, representada neste acto pelo Sr. {{ representative_name }}, na qualidade de Gerente, com poderes para o acto, doravante designado, abreviadamente "SENHORIO".'),
    ('para', '\nE\n'),
    ('para', '{{ inquilino }}, pessoa singular, portador do {{ document_type }} {{ document_number }}, emitido em {{ document_issue_date }}, válido até {{ document_expiry_date }}, NIF {{ inquilino_nif }}, adiante abreviadamente designado por ARRENDATÁRIO.'),
    ('para', '\nIndividualmente designadas por "Parte" e colectivamente por "Partes"'),
    ('para', ''),
    ('para', 'CONSIDERANDO QUE:'),
    ('para', 'O SENHORIO declara ser o dono e legítimo proprietário do {{ endereco_imovel }}, na província de Luanda, destinado a habitação, doravante designado "imóvel".'),
    ('para', ''),
    ('para', 'CLÁUSULA SEGUNDA (Vigência)'),
    ('para', 'O contrato terá o prazo de 1(um) ano com início a {{ start_date_written }} e término a {{ end_date_written }}, sendo automaticamente renovável por períodos sucessivos.'),
    ('para', ''),
    ('para', 'CLÁUSULA TERCEIRA (Renda e Formas de Pagamento)'),
    ('para', 'As Partes acordam um valor mensal da renda equivalente a {{ valor_renda }}.'),
    ('para', 'O pagamento da renda será efectuado por {{ forma_pagamento }} nos primeiros 8 dias úteis.'),
    ('para', 'Caução: {{ valor_caucao }}.'),
    ('para', 'Taxa de condomínio: {{ taxa_condominio }}.'),
    ('para', ''),
    ('para', 'CLÁUSULA DÉCIMA TERCEIRA (Notificações)'),
    ('para', 'Os contactos do Arrendatário:'),
    ('para', 'Tel.: {{ inquilino_contact }}'),
    ('para', 'Email: {{ inquilino_email }}'),
    ('para', ''),
    ('para', 'Luanda, aos {{ contract_date_local }}'),
    ('para', ''),
    ('para', 'P/SENHORIO __________________________'),
    ('para', 'P/ ARRENDATÁRIO __________________________')
]


def create(path):
    doc = Document()
    for kind, text in CONTRACT_TEXT:
        if kind == 'heading':
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)
    doc.save(path)

if __name__ == '__main__':
    create('contract_template.docx')
    print('Wrote contract_template.docx with clean placeholders')
